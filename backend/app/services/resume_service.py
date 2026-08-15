import io
import re
import json
from typing import Dict, Any, List, Optional, Tuple
import pypdf
import docx
from app.models.candidate import Candidate
from app.models.student_profile import StudentProfile
from app.models.resume import Resume

# ==============================================================================
# UNIVERSAL MULTI-DOMAIN SKILL TAXONOMY (Exact Word Boundary Matchers)
# ==============================================================================
SKILL_CATEGORIES: Dict[str, List[Tuple[str, str]]] = {
    "Programming Languages": [
        ("Python", r"(?i)\b(?:python(?:\s*3)?|py)\b"),
        ("Java", r"(?i)\b(?:core\s+java|java(?:\s*(?:8|11|17|21))?)\b(?!\s*script)"),
        ("JavaScript", r"(?i)\b(?:javascript|js|es6)\b"),
        ("TypeScript", r"(?i)\b(?:typescript|ts)\b"),
        ("C++", r"(?i)(?<![a-zA-Z0-9])c\+\+(?![a-zA-Z0-9+])"),
        ("C#", r"(?i)(?:(?<![a-zA-Z0-9])c#(?![a-zA-Z0-9#])|\bc\s*sharp\b)"),
        ("Go", r"(?i)\b(?:golang|go\s+language)\b"),
        ("Rust", r"(?i)\brust\b"),
        ("PHP", r"(?i)\bphp\b"),
        ("Ruby", r"(?i)\bruby\b"),
        ("Kotlin", r"(?i)\bkotlin\b"),
        ("Swift", r"(?i)\bswift\b"),
        ("Scala", r"(?i)\bscala\b"),
        ("R", r"(?i)\b(?:r\s+programming|r\s+language)\b"),
        ("SQL", r"(?i)\bsql\b"),
        ("HTML5/CSS3", r"(?i)\b(?:html5?|css3?)\b"),
        ("Bash/Shell", r"(?i)\b(?:bash|shell\s+scripting|powershell)\b")
    ],
    "Frontend & Web Frameworks": [
        ("React.js", r"(?i)\b(?:react|react\.js|reactjs)\b"),
        ("Angular", r"(?i)\bangular(?:\s*(?:2|4|6|8|10|12|14|16|17))?\b"),
        ("Vue.js", r"(?i)\b(?:vue|vue\.js|vuejs)\b"),
        ("Next.js", r"(?i)\bnext\.js\b"),
        ("Tailwind CSS", r"(?i)\btailwind(?:\s*css)?\b"),
        ("Bootstrap", r"(?i)\bbootstrap\b"),
        ("Redux", r"(?i)\bredux\b"),
        ("GraphQL", r"(?i)\bgraphql\b"),
        ("Webpack/Vite", r"(?i)\b(?:webpack|vite)\b")
    ],
    "Backend & API Frameworks": [
        ("Node.js", r"(?i)\b(?:node|node\.js|nodejs)\b"),
        ("Express.js", r"(?i)\bexpress(?:\.js)?\b"),
        ("FastAPI", r"(?i)\bfastapi\b"),
        ("Django", r"(?i)\bdjango\b"),
        ("Flask", r"(?i)\bflask\b"),
        ("Spring Boot", r"(?i)\bspring\s+boot\b"),
        ("ASP.NET", r"(?i)\b(?:asp\.net|\.net\s+core)\b"),
        ("REST API", r"(?i)\b(?:rest\s*api|restful\s*api|rest\s*web\s*services)\b"),
        ("Microservices", r"(?i)\bmicroservices?\b")
    ],
    "Cloud & Infrastructure": [
        ("Amazon Web Services (AWS)", r"(?i)\b(?:aws|amazon\s+web\s+services|ec2|s3|lambda)\b"),
        ("Microsoft Azure", r"(?i)\b(?:azure|microsoft\s+azure)\b"),
        ("Google Cloud (GCP)", r"(?i)\b(?:gcp|google\s+cloud)\b"),
        ("Docker", r"(?i)\bdocker\b"),
        ("Kubernetes", r"(?i)\b(?:kubernetes|k8s)\b"),
        ("Terraform", r"(?i)\bterraform\b"),
        ("Linux", r"(?i)\blinux\b"),
        ("Serverless", r"(?i)\bserverless\b")
    ],
    "DevOps & CI/CD": [
        ("Jenkins", r"(?i)\bjenkins\b"),
        ("GitLab CI", r"(?i)\bgitlab(?:\s+ci)?\b"),
        ("GitHub Actions", r"(?i)\bgithub\s+actions\b"),
        ("CI/CD Pipelines", r"(?i)\b(?:ci/cd|cicd|continuous\s+integration)\b"),
        ("Ansible", r"(?i)\bansible\b"),
        ("Prometheus", r"(?i)\bprometheus\b"),
        ("Grafana", r"(?i)\bgrafana\b")
    ],
    "Databases & Storage": [
        ("PostgreSQL", r"(?i)\b(?:postgresql|postgres)\b"),
        ("MySQL", r"(?i)\bmysql\b"),
        ("MongoDB", r"(?i)\bmongodb\b"),
        ("Redis", r"(?i)\bredis\b"),
        ("Oracle Database", r"(?i)\boracle(?:\s+database|\s+db)?\b"),
        ("Elasticsearch", r"(?i)\belasticsearch\b"),
        ("Snowflake", r"(?i)\bsnowflake\b")
    ],
    "Data Analytics, AI & BI": [
        ("Power BI", r"(?i)\bpower\s*bi\b"),
        ("Tableau", r"(?i)\btableau\b"),
        ("Advanced Excel", r"(?i)\b(?:advanced\s+excel|ms\s+excel|vlookup|pivots?)\b"),
        ("Pandas & NumPy", r"(?i)\b(?:pandas|numpy)\b"),
        ("Machine Learning", r"(?i)\b(?:machine\s+learning|scikit-learn|sklearn)\b"),
        ("Deep Learning / PyTorch", r"(?i)\b(?:pytorch|tensorflow|keras|deep\s+learning)\b"),
        ("ETL Pipelines", r"(?i)\betl\b"),
        ("Data Warehousing", r"(?i)\bdata\s+warehousing\b")
    ],
    "Business Analysis & Product": [
        ("Business Analysis", r"(?i)\bbusiness\s+analysis\b"),
        ("Requirements Gathering", r"(?i)\b(?:requirements\s+gathering|brd|frd)\b"),
        ("User Stories & Backlog", r"(?i)\b(?:user\s+stories|product\s+backlog)\b"),
        ("Stakeholder Management", r"(?i)\bstakeholder\s+management\b"),
        ("Agile & Scrum", r"(?i)\b(?:agile|scrum|kanban|sprint\s+planning)\b"),
        ("Jira & Confluence", r"(?i)\b(?:jira|confluence)\b")
    ],
    "Quality Assurance & Automation": [
        ("Selenium WebDriver", r"(?i)\bselenium(?:\s+webdriver)?\b"),
        ("Test Automation", r"(?i)\b(?:test\s+automation|automation\s+testing)\b"),
        ("Postman", r"(?i)\bpostman\b"),
        ("Playwright", r"(?i)\bplaywright\b"),
        ("Cypress", r"(?i)\bcypress\b"),
        ("TestNG & JUnit", r"(?i)\b(?:testng|junit)\b"),
        ("Pytest", r"(?i)\bpytest\b"),
        ("Manual Testing", r"(?i)\bmanual\s+testing\b"),
        ("Performance Testing / JMeter", r"(?i)\b(?:jmeter|performance\s+testing)\b"),
        ("Defect Management", r"(?i)\b(?:defect|bug)\s+(?:tracking|lifecycle|management)\b")
    ],
    "Cybersecurity": [
        ("Penetration Testing", r"(?i)\b(?:penetration\s+testing|pen\s+testing)\b"),
        ("Vulnerability Assessment", r"(?i)\bvulnerability\s+assessment\b"),
        ("SOC / SIEM", r"(?i)\b(?:soc|siem|splunk)\b"),
        ("Network Security", r"(?i)\bnetwork\s+security\b"),
        ("OWASP Security", r"(?i)\bowasp\b")
    ],
    "UI/UX Design": [
        ("Figma", r"(?i)\bfigma\b"),
        ("UI/UX Design", r"(?i)\b(?:ui/ux|user\s+interface|user\s+experience)\b"),
        ("Wireframing & Prototyping", r"(?i)\b(?:wireframing|prototyping)\b"),
        ("Design Systems", r"(?i)\bdesign\s+systems?\b")
    ]
}

DEGREE_MATCHERS = [
    ("Bachelor of Technology (B.Tech)", r"(?i)\b(?:bachelors?\s+of\s+technology|b\.?\s*tech(?:nology)?)\b"),
    ("Bachelor of Engineering (B.E.)", r"(?i)\b(?:bachelors?\s+of\s+engineering|b\.?\s*e\.?)\b"),
    ("Diploma", r"(?i)\bdiploma\b"),
    ("Bachelor of Computer Applications (BCA)", r"(?i)\b(?:bca|bachelors?\s+of\s+computer\s+applications)\b"),
    ("Bachelor of Science (B.Sc)", r"(?i)\b(?:b\.?\s*sc(?:ience)?|bachelors?\s+of\s+science)\b"),
    ("Master of Computer Applications (MCA)", r"(?i)\b(?:mca|masters?\s+of\s+computer\s+applications)\b"),
    ("Master of Technology (M.Tech)", r"(?i)\b(?:m\.?\s*tech(?:nology)?|masters?\s+of\s+technology)\b"),
    ("Master of Engineering (M.E.)", r"(?i)\b(?:m\.?\s*e\.?|masters?\s+of\s+engineering)\b"),
    ("Master of Science (M.Sc)", r"(?i)\b(?:m\.?\s*sc(?:ience)?|masters?\s+of\s+science)\b"),
    ("Master of Business Administration (MBA)", r"(?i)\b(?:mba|masters?\s+of\s+business\s+administration)\b"),
    ("Bachelor of Commerce (B.Com)", r"(?i)\b(?:b\.?\s*com|bachelors?\s+of\s+commerce)\b"),
    ("Bachelor of Business Administration (BBA)", r"(?i)\b(?:bba|bachelors?\s+of\s+business\s+administration)\b"),
    ("Higher Secondary / 12th Grade", r"(?i)\b(?:12th\s+(?:grade|standard|pass)|higher\s+secondary|hsc)\b"),
    ("Secondary School / 10th Grade", r"(?i)\b(?:10th\s+(?:grade|standard|pass)|ssc)\b"),
]

def clean_and_normalize_text(raw_text: str) -> str:
    """
    Sanitizes extracted text:
    - Removes null bytes and control characters
    - Strips PDF headers, binary markers, and xref/trailer artifacts
    - Normalizes unicode spaces, dashes, and quotes
    - Preserves meaningful section line breaks
    """
    if not raw_text:
        return ""

    text = raw_text.replace("\x00", "")
    text = re.sub(r"[\u00a0\u1680\u2000-\u200b\u202f\u205f\u3000]", " ", text)
    text = re.sub(r"[\u2010-\u2015\u2212\ufffd\x96\x97]", "-", text)
    text = re.sub(r"[\u2018\u2019\u201a\u201b]", "'", text)
    text = re.sub(r"[\u201c\u201d\u201e\u201f]", '"', text)

    cleaned_lines = []
    for line in text.splitlines():
        line_str = line.strip()
        if not line_str:
            cleaned_lines.append("")
            continue
        if re.match(r"^%PDF-\d+\.\d+", line_str, re.IGNORECASE):
            continue
        if re.match(r"^\d+\s+\d+\s+obj\b", line_str, re.IGNORECASE):
            continue
        if line_str.lower() in ["endobj", "xref", "trailer", "startxref", "stream", "endstream"]:
            continue
        if re.match(r"^/Type\s*/", line_str):
            continue
        printable_ratio = sum(1 for c in line_str if 32 <= ord(c) <= 126 or ord(c) in [9, 10, 13]) / max(len(line_str), 1)
        if printable_ratio < 0.6 and len(line_str) > 10:
            continue
        cleaned_lines.append(line_str)

    result = "\n".join(cleaned_lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()

def extract_text_from_file_bytes(content: bytes, file_name: str) -> str:
    """
    Extracts clean, readable text from uploaded resume bytes (.pdf, .docx, .txt).
    """
    if not content or len(content) == 0:
        return ""

    ext = file_name.lower().split(".")[-1] if "." in file_name else ""

    if ext == "pdf":
        try:
            reader = pypdf.PdfReader(io.BytesIO(content))
            page_texts = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    page_texts.append(extracted.strip())
            full_pdf_text = "\n\n".join(page_texts)
            return clean_and_normalize_text(full_pdf_text)
        except Exception:
            return ""

    elif ext == "docx":
        try:
            doc = docx.Document(io.BytesIO(content))
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text_parts.append(" | ".join(row_data))
            return clean_and_normalize_text("\n".join(text_parts))
        except Exception:
            return ""

    else:
        for encoding in ["utf-8-sig", "utf-8", "cp1252", "latin-1"]:
            try:
                decoded = content.decode(encoding)
                return clean_and_normalize_text(decoded)
            except (UnicodeDecodeError, LookupError):
                continue
        return clean_and_normalize_text(content.decode("utf-8", errors="ignore"))

def extract_education_entries(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    text = clean_and_normalize_text(text)
    entries = []
    lines = text.splitlines()

    found_matches = []
    last_line_idx = -10
    for i, line in enumerate(lines):
        clean_l = line.strip()
        if not clean_l:
            continue
        if i - last_line_idx <= 2 and any(k in clean_l.lower() for k in ["polytechnic", "board", "university", "college", "campus"]):
            continue
        for deg_name, deg_pat in DEGREE_MATCHERS:
            m = re.search(deg_pat, clean_l)
            if m:
                found_matches.append((i, deg_name, clean_l, m))
                last_line_idx = i
                break

    for idx, (line_idx, deg_name, matched_line, match_obj) in enumerate(found_matches):
        end_idx = found_matches[idx + 1][0] if idx + 1 < len(found_matches) else min(len(lines), line_idx + 7)
        block_lines = [lines[j].strip() for j in range(line_idx, end_idx) if lines[j].strip()]
        block_text = " \n ".join(block_lines)

        spec = None
        after_deg = matched_line[match_obj.end():].strip()
        clean_after = re.sub(r"^[^a-zA-Z0-9]+", "", after_deg).strip()
        clean_after = re.sub(r"^(?:in\b|in the field of\b|of\b)\s*", "", clean_after, flags=re.IGNORECASE).strip()
        clean_after = re.sub(r"[\)\]]+$", "", clean_after).strip()
        if clean_after and len(clean_after) > 2 and not any(k in clean_after.lower() for k in ["university", "college", "campus", "institute", "board", "cgpa", "gpa", "202", "201", "199"]):
            spec = clean_after
        else:
            spec_match = re.search(r"(?i)(?:[-–—:]|\bin\b)\s*([A-Za-z\s&,]+(?:engineering|technology|science|applications|management|commerce|arts)?)", matched_line)
            if spec_match:
                cand_spec = spec_match.group(1).strip()
                if len(cand_spec) > 2 and not any(k in cand_spec.lower() for k in ["university", "college", "campus", "institute", "board"]):
                    spec = cand_spec

        inst = None
        for b_line in block_lines[1:]:
            if re.search(r"(?i)(?:university|college|campus|institute|polytechnic|board|school|academy|vidyalaya)", b_line):
                inst = b_line
                break
        if not inst and len(block_lines) > 1:
            for b_line in block_lines[1:]:
                if not re.search(r"(?i)\b(?:cgpa|percentage|gpa|%\b|\b20\d\d\b)", b_line):
                    inst = b_line
                    break

        year = None
        year_match = re.search(r"\b(20\d\d|19\d\d)\b", block_text)
        if year_match:
            year = int(year_match.group(1))

        score = None
        score_match = re.search(r"(?i)(?:cgpa\s*[:\-]?\s*(\d+(?:\.\d+)?)|(\d+(?:\.\d+)?\s*%|\d+(?:\.\d+)?\s*cgpa)|gpa\s*[:\-]?\s*(\d+(?:\.\d+)?(?:\s*\/\s*\d+)?))", block_text)
        if score_match:
            score = score_match.group(0).strip()

        display_parts = [deg_name]
        if spec:
            display_parts[0] = f"{deg_name} — {spec}"
        if inst:
            display_parts.append(inst)
        if year:
            display_parts.append(str(year))
        if score:
            display_parts.append(score)

        display_text = " | ".join(display_parts)
        entries.append({
            "degree": deg_name,
            "specialization": spec,
            "institution": inst,
            "year": year,
            "score": score,
            "display_text": display_text
        })

    return entries

def detect_career_domain_and_roles(text: str, categorized_skills: Dict[str, List[str]]) -> Dict[str, Any]:
    """
    Holistically detects the primary career domain and likely professional roles
    without fabricating certainty or forcing QA specialization.
    """
    text_lower = text.lower()

    domain_scores = {
        "Software Development": 0,
        "DevOps & Cloud Engineering": 0,
        "Data Analytics & AI": 0,
        "Business Analysis & Product Management": 0,
        "Quality Assurance & Test Automation": 0,
        "Cybersecurity": 0,
        "UI/UX & Product Design": 0
    }

    # 1. Skill category signals
    if categorized_skills.get("Programming Languages"):
        domain_scores["Software Development"] += len(categorized_skills["Programming Languages"]) * 3
    if categorized_skills.get("Frontend & Web Frameworks"):
        domain_scores["Software Development"] += len(categorized_skills["Frontend & Web Frameworks"]) * 4
    if categorized_skills.get("Backend & API Frameworks"):
        domain_scores["Software Development"] += len(categorized_skills["Backend & API Frameworks"]) * 4
    if categorized_skills.get("Cloud & Infrastructure"):
        domain_scores["DevOps & Cloud Engineering"] += len(categorized_skills["Cloud & Infrastructure"]) * 4
    if categorized_skills.get("DevOps & CI/CD"):
        domain_scores["DevOps & Cloud Engineering"] += len(categorized_skills["DevOps & CI/CD"]) * 4
    if categorized_skills.get("Data Analytics, AI & BI"):
        domain_scores["Data Analytics & AI"] += len(categorized_skills["Data Analytics, AI & BI"]) * 4
    if categorized_skills.get("Business Analysis & Product"):
        domain_scores["Business Analysis & Product Management"] += len(categorized_skills["Business Analysis & Product"]) * 4
    if categorized_skills.get("Quality Assurance & Automation"):
        domain_scores["Quality Assurance & Test Automation"] += len(categorized_skills["Quality Assurance & Automation"]) * 4
    if categorized_skills.get("Cybersecurity"):
        domain_scores["Cybersecurity"] += len(categorized_skills["Cybersecurity"]) * 4
    if categorized_skills.get("UI/UX Design"):
        domain_scores["UI/UX & Product Design"] += len(categorized_skills["UI/UX Design"]) * 4

    # 2. Text keyword signals
    if re.search(r"(?i)\b(?:full\s*stack|backend\s*developer|frontend\s*developer|software\s*engineer|software\s*developer)\b", text_lower):
        domain_scores["Software Development"] += 8
    if re.search(r"(?i)\b(?:devops|site\s*reliability|sre|cloud\s*engineer|infrastructure)\b", text_lower):
        domain_scores["DevOps & Cloud Engineering"] += 8
    if re.search(r"(?i)\b(?:data\s*analyst|bi\s*analyst|data\s*scientist|machine\s*learning|business\s*intelligence)\b", text_lower):
        domain_scores["Data Analytics & AI"] += 8
    if re.search(r"(?i)\b(?:business\s*analyst|product\s*manager|scrum\s*master|agile\s*coach)\b", text_lower):
        domain_scores["Business Analysis & Product Management"] += 8
    if re.search(r"(?i)\b(?:qa\s*engineer|test\s*automation|sdet|manual\s*tester|quality\s*assurance)\b", text_lower):
        domain_scores["Quality Assurance & Test Automation"] += 8
    if re.search(r"(?i)\b(?:security\s*analyst|soc\s*analyst|penetration\s*testing|cybersecurity)\b", text_lower):
        domain_scores["Cybersecurity"] += 8
    if re.search(r"(?i)\b(?:ui/ux|product\s*designer|ux\s*researcher|interaction\s*designer)\b", text_lower):
        domain_scores["UI/UX & Product Design"] += 8

    # Pick top domain
    sorted_domains = sorted(domain_scores.items(), key=lambda x: x[1], reverse=True)
    top_domain, top_score = sorted_domains[0]

    likely_roles = []
    explanation = ""

    if top_score < 4:
        return {
            "career_domain": "General Technical Profile",
            "likely_roles": ["Software Associate", "Technical Specialist"],
            "domain_explanation": "Insufficient specific specialization detected; evaluated as a foundational professional profile."
        }

    if top_domain == "Software Development":
        if "React.js" in text or "Angular" in text or "Frontend & Web Frameworks" in categorized_skills:
            likely_roles = ["Full Stack Developer", "Frontend Developer", "Web Applications Engineer"]
        elif "Python" in text or "FastAPI" in text or "Django" in text:
            likely_roles = ["Backend Developer", "Python Developer", "API Engineer"]
        elif "Java" in text or "Spring Boot" in text:
            likely_roles = ["Java Developer", "Backend Software Engineer", "Enterprise Application Developer"]
        else:
            likely_roles = ["Software Engineer", "Full Stack Developer", "Application Developer"]
        explanation = f"Detected primary focus in Software Development based on core programming and modern application framework skills."

    elif top_domain == "DevOps & Cloud Engineering":
        likely_roles = ["DevOps Engineer", "Cloud Infrastructure Engineer", "Site Reliability Engineer (SRE)"]
        explanation = f"Detected primary focus in DevOps & Cloud Engineering based on containerization, CI/CD pipelines, and cloud platform tools."

    elif top_domain == "Data Analytics & AI":
        likely_roles = ["Data Analyst", "Business Intelligence (BI) Analyst", "Data Engineer"]
        explanation = f"Detected primary focus in Data Analytics & BI based on SQL, visualization tools, and analytical data modeling."

    elif top_domain == "Quality Assurance & Test Automation":
        likely_roles = ["QA Automation Engineer", "Software Development Engineer in Test (SDET)", "QA Analyst"]
        explanation = f"Detected primary focus in Quality Assurance based on test automation frameworks, API validation, and quality lifecycle metrics."

    elif top_domain == "Business Analysis & Product Management":
        likely_roles = ["Business Analyst", "Product Analyst", "Associate Product Manager"]
        explanation = f"Detected primary focus in Business Analysis & Product Management based on requirement specifications, user stories, and agile delivery."

    elif top_domain == "Cybersecurity":
        likely_roles = ["Information Security Analyst", "SOC Analyst", "Security Engineer"]
        explanation = f"Detected primary focus in Cybersecurity based on security monitoring, vulnerability assessment, and threat mitigation."

    elif top_domain == "UI/UX & Product Design":
        likely_roles = ["UI/UX Designer", "Product Designer", "User Experience Researcher"]
        explanation = f"Detected primary focus in UI/UX Design based on interface wireframing, design systems, and user research."

    return {
        "career_domain": top_domain,
        "likely_roles": likely_roles,
        "domain_explanation": explanation
    }

def parse_resume_text(text: str) -> Dict[str, Any]:
    text = clean_and_normalize_text(text)
    if not text or len(text.strip()) < 10:
        return {
            "full_name": None,
            "email": None,
            "phone": None,
            "skills": [],
            "categorized_skills": {},
            "education": [],
            "education_entries": [],
            "career_domain": "Unspecified",
            "likely_roles": [],
            "domain_explanation": "No readable text detected.",
            "linkedin_url": None,
            "github_url": None,
            "raw_text_snippet": "Document does not contain readable text.",
            "extraction_warning": "No readable text could be extracted from this file."
        }

    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # 1. Extract Email
    email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,7}\b", text)
    email = email_match.group(0).strip() if email_match else None

    # 2. Extract Phone
    phone = None
    in_phone = re.search(r"(?:\+91[\s-]?)?[6-9]\d{4}[\s-]?\d{5}\b", text)
    if in_phone:
        phone = in_phone.group(0).strip()
    else:
        intl_phone = re.search(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b", text)
        if intl_phone:
            phone = intl_phone.group(0).strip()

    # 3. Extract Name
    candidate_name = None
    name_label_match = re.search(r"(?i)\b(?:name|full\s*name|candidate\s*name)\s*[:\-]\s*([A-Za-z\s\.\'\-]{3,35})", text)
    if name_label_match:
        cand = name_label_match.group(1).strip()
        if cand and not any(k in cand.lower() for k in ["resume", "curriculum", "email", "phone", "profile", "summary"]):
            candidate_name = cand

    if not candidate_name and lines:
        for line in lines[:5]:
            clean_l = line.strip()
            if "@" in clean_l or "http" in clean_l or ".com" in clean_l or re.search(r"\d{4}", clean_l):
                continue
            lower_l = clean_l.lower()
            if any(k in lower_l for k in [
                "resume", "curriculum", "vitae", "cv", "page", "contact", "summary",
                "experience", "education", "skills", "projects", "objective", "profile",
                "phone", "email", "address", "declaration", "about me", "details"
            ]):
                continue
            words = clean_l.split()
            if 2 <= len(words) <= 4 and all(re.match(r"^[A-Z][a-zA-Z\.\'\-]*$", w) for w in words):
                if len(clean_l) <= 35:
                    candidate_name = clean_l
                    break

    if not candidate_name and email:
        local_part = email.split("@")[0]
        if "." in local_part:
            parts = [p.capitalize() for p in local_part.split(".") if p.isalpha() and len(p) > 1]
            if 2 <= len(parts) <= 3:
                candidate_name = " ".join(parts)

    # 4. Extract Structured Categorized Skills
    categorized_skills: Dict[str, List[str]] = {}
    flat_skills_set = set()

    for category, skill_list in SKILL_CATEGORIES.items():
        matched_in_cat = []
        for skill_name, pattern in skill_list:
            if re.search(pattern, text):
                matched_in_cat.append(skill_name)
                flat_skills_set.add(skill_name)
        if matched_in_cat:
            categorized_skills[category] = matched_in_cat

    # 5. Extract Multi-entry Education
    edu_entries = extract_education_entries(text)
    edu_display_list = [e["display_text"] for e in edu_entries]

    # 6. Extract Professional Links
    linkedin_match = re.search(r"(?i)\bhttps?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_\-\.\/]+", text)
    linkedin_url = linkedin_match.group(0).rstrip(".,;/ ") if linkedin_match else None

    github_match = re.search(r"(?i)\bhttps?://(?:www\.)?github\.com/[a-zA-Z0-9_\-\.\/]+", text)
    github_url = github_match.group(0).rstrip(".,;/ ") if github_match else None

    # 7. Detect Career Domain & Roles
    domain_info = detect_career_domain_and_roles(text, categorized_skills)

    clean_snippet = re.sub(r"\s+", " ", text[:500]).strip()

    return {
        "full_name": candidate_name,
        "email": email,
        "phone": phone,
        "skills": sorted(list(flat_skills_set)),
        "categorized_skills": categorized_skills,
        "education": edu_display_list,
        "education_entries": edu_entries,
        "career_domain": domain_info["career_domain"],
        "likely_roles": domain_info["likely_roles"],
        "domain_explanation": domain_info["domain_explanation"],
        "linkedin_url": linkedin_url,
        "github_url": github_url,
        "raw_text_snippet": clean_snippet,
        "extraction_warning": None
    }

def calculate_ats_score(text: str, parsed: Dict[str, Any]) -> Dict[str, Any]:
    """
    Calculates the domain-adaptive NxtMov ATS Score (0-100) using a transparent weighted scoring model:
    1. Contact Information (10%)
    2. Resume Structure (15%)
    3. Keyword Coverage & Role Relevance (20%)
    4. Experience & Projects Quality (20%)
    5. Skills Diversity across Modern Toolchains (15%)
    6. Education Multi-Tier (10%)
    7. ATS Readability & Formatting (10%)
    """
    breakdown = {
        "contact_info": 0,
        "structure": 0,
        "keyword_coverage": 0,
        "experience_projects": 0,
        "skills_diversity": 0,
        "education": 0,
        "ats_readability": 0
    }
    strengths = []
    improvements = []
    warnings = []

    text_lower = text.lower()
    words = text.split()

    # 1. Contact Info (10%)
    c_score = 0
    name = parsed.get("full_name")
    if name and len(name.strip()) > 3 and not any(ch.isdigit() for ch in name):
        c_score += 3
    email = parsed.get("email")
    if email and re.match(r"^[\w\.-]+@[\w\.-]+\.\w+$", email):
        c_score += 3
    phone = parsed.get("phone")
    if phone and len(re.sub(r"\D", "", phone)) >= 10:
        c_score += 2
    if parsed.get("linkedin_url") or parsed.get("github_url"):
        c_score += 2
    else:
        warnings.append("• No professional LinkedIn or GitHub profile link detected in contact section")

    breakdown["contact_info"] = min(10, c_score)
    if c_score >= 8:
        strengths.append("✓ Complete verified contact profile (Name, Email & Phone detected)")
    else:
        improvements.append("• Provide complete verified contact details including LinkedIn and GitHub")

    # 2. Resume Structure (15%)
    s_score = 0
    if any(k in text_lower for k in ["summary", "objective", "profile", "about"]):
        s_score += 3
    if any(k in text_lower for k in ["skill", "technologies", "technical", "competencies"]):
        s_score += 3
    if any(k in text_lower for k in ["experience", "employment", "work history", "projects"]):
        s_score += 3
    if any(k in text_lower for k in ["education", "academic", "qualification"]):
        s_score += 3
    if len(text.splitlines()) >= 15:
        s_score += 3
    breakdown["structure"] = min(15, s_score)

    if s_score >= 12:
        strengths.append("✓ Standard ATS-compliant section hierarchy detected")
    else:
        improvements.append("• Structure resume with clear standard section headings (Summary, Skills, Experience, Education)")

    # 3. Keyword Coverage & Role Relevance (20%)
    skill_count = len(parsed.get("skills", []))
    domain = parsed.get("career_domain", "General")
    if skill_count >= 12:
        kw_score = 20
        strengths.append(f"✓ Strong industry keyword density ({skill_count} relevant technical skills in {domain})")
    elif skill_count >= 8:
        kw_score = 16
        strengths.append(f"✓ Good keyword coverage ({skill_count} technical skills)")
    elif skill_count >= 4:
        kw_score = 12
        improvements.append("• Expand industry keywords and framework coverage for higher ATS matching")
    elif skill_count >= 2:
        kw_score = 8
        improvements.append("• Include core programming languages, frameworks, and domain tooling")
    else:
        kw_score = 4
        improvements.append("• Add a dedicated Technical Skills section with role-specific keywords")
    breakdown["keyword_coverage"] = kw_score

    # 4. Experience & Projects Quality (20%)
    exp_score = 0
    action_verbs = [
        "developed", "tested", "designed", "implemented", "automated", "executed",
        "architected", "deployed", "optimized", "engineered", "built", "spearheaded",
        "managed", "delivered", "configured", "created", "integrated", "analyzed"
    ]
    found_verbs = [v for v in action_verbs if v in text_lower]
    if len(found_verbs) >= 4:
        exp_score += 8
        strengths.append("✓ Strong action verbs used throughout experience descriptions")
    elif len(found_verbs) >= 2:
        exp_score += 5
    else:
        exp_score += 2
        improvements.append("• Begin experience bullet points with strong action verbs (e.g. Engineered, Deployed, Automated)")

    if any(k in text_lower for k in ["project", "framework", "architecture", "system", "application", "pipeline"]):
        exp_score += 4

    # Detect quantified metrics
    metrics_match = re.findall(r"\b\d+%\b|\b\d+\s+bugs?\b|\b\d+\s+test cases?\b|\b\d+\s+years?\b|\b\d+x\b|\b\d+\s+ms\b|\b\d+\s+users\b", text_lower)
    if len(metrics_match) >= 2:
        exp_score += 8
        strengths.append("✓ Quantified impact & measurable achievements detected")
    elif len(metrics_match) == 1:
        exp_score += 5
        strengths.append("✓ Measurable metrics detected in project experience")
    else:
        exp_score += 2
        improvements.append("• Add quantified metrics and measurable achievements (e.g. % performance increase, scale, test coverage)")

    # Detect weak phrases and warn
    weak_phrases = ["worked on", "responsible for", "handled", "worked with"]
    found_weak = [p for p in weak_phrases if p in text_lower]
    if found_weak:
        warnings.append("• Some experience bullets use passive/responsibility phrases ('" + "', '".join(found_weak[:2]) + "') rather than achievement-focused action verbs")

    breakdown["experience_projects"] = min(20, exp_score)

    # 5. Skills Diversity (15%)
    cat_count = len(parsed.get("categorized_skills", {}))
    if cat_count >= 4:
        sk_score = 15
        strengths.append(f"✓ Balanced skill diversity across {cat_count} technical categories")
    elif cat_count >= 2:
        sk_score = 11
    elif cat_count >= 1:
        sk_score = 7
        improvements.append("• Diversify skills across backend, cloud, databases, and CI/CD tools")
    else:
        sk_score = 3
    breakdown["skills_diversity"] = sk_score

    # 6. Education (10%)
    edu_entries = parsed.get("education_entries", [])
    if len(edu_entries) >= 2:
        edu_score = 10
        strengths.append(f"✓ Multi-tiered academic qualifications verified ({len(edu_entries)} degrees/diplomas)")
    elif len(edu_entries) == 1:
        edu_score = 8
        strengths.append(f"✓ Academic qualification verified ({edu_entries[0].get('degree', 'Degree')})")
    elif parsed.get("education"):
        edu_score = 6
    else:
        edu_score = 2
        improvements.append("• Include degree, university/board, graduation year, and GPA/Percentage")
    breakdown["education"] = edu_score

    # 7. ATS Readability (10%)
    if len(words) >= 150:
        read_score = 10
        strengths.append("✓ Clean text parsing and optimal length for single/two-page ATS parsing")
    elif len(words) >= 60:
        read_score = 6
        improvements.append("• Elaborate work history, project descriptions, and responsibilities")
    else:
        read_score = 3
        improvements.append("• Upload a complete resume document with comprehensive details")
        warnings.append("• Resume length is very short (<60 words) which may limit ATS keyword indexation")

    breakdown["ats_readability"] = read_score

    total_score = sum(breakdown.values())
    total_score = max(20, min(100, total_score))

    return {
        "ats_score": total_score,
        "quality_score": total_score,
        "career_domain": parsed.get("career_domain"),
        "likely_roles": parsed.get("likely_roles", []),
        "domain_explanation": parsed.get("domain_explanation"),
        "score_breakdown": breakdown,
        "strengths": strengths,
        "improvements": improvements,
        "warnings": warnings
    }

calculate_resume_quality_score = calculate_ats_score

def calculate_profile_completeness(cand: Candidate, profile: Optional[StudentProfile] = None) -> Dict[str, Any]:
    score = 0
    missing = []

    # 1. Identity & Contact (30%)
    if cand.full_name:
        score += 10
    else:
        missing.append("Full Name")

    if cand.email:
        score += 10
    else:
        missing.append("Email Address")

    if cand.phone or (profile and profile.city):
        score += 10
    else:
        missing.append("Phone Number / Location")

    # 2. Headline & Career Preferences (20%)
    if profile and profile.headline:
        score += 10
    else:
        missing.append("Professional Headline")

    if profile and (profile.preferred_roles or profile.career_objective):
        score += 10
    else:
        missing.append("Career Preferences / Preferred Roles")

    # 3. Skills (25%)
    if cand.primary_skills or (profile and (profile.programming_languages or profile.frameworks)):
        score += 25
    else:
        missing.append("Key Skills & Technologies")

    # 4. Education (15%)
    if profile and (profile.degree or profile.highest_qualification or profile.college_university):
        score += 15
    else:
        missing.append("Education Qualifications")

    # 5. Links & Resume (10%)
    if cand.resume_url or (profile and (profile.linkedin_url or profile.github_url)):
        score += 10
    else:
        missing.append("Resume Upload / Online Profiles")

    score = min(100, score)
    return {
        "completeness_score": score,
        "missing_sections": missing
    }
