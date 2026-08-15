import io
import pytest
from fastapi.testclient import TestClient
import docx
from app.main import app
from app.core.security import create_access_token

client = TestClient(app)


def get_authenticated_headers():
    token_resp = client.post(
        "/api/v1/auth/login",
        data={"username": "demo@nxtmov.local", "password": "NxtMov@123"}
    )
    assert token_resp.status_code == 200
    token = token_resp.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def test_resume_upload_unauthenticated_fails_401():
    file_bytes = b"%PDF-1.4\nSome sample resume text here for test"
    files = {"file": ("resume.pdf", io.BytesIO(file_bytes), "application/pdf")}
    res = client.post("/api/v1/resumes/upload", files=files)
    assert res.status_code == 401
    assert "detail" in res.json()


def test_resume_upload_empty_file_fails_400():
    headers = get_authenticated_headers()
    files = {"file": ("empty.pdf", io.BytesIO(b""), "application/pdf")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 400
    assert "empty" in res.json()["detail"].lower()


def test_resume_upload_unsupported_format_fails_415():
    headers = get_authenticated_headers()
    files = {"file": ("image.png", io.BytesIO(b"\x89PNG\r\n\x1a\n"), "image/png")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 415
    assert "unsupported" in res.json()["detail"].lower()


def test_resume_upload_oversized_file_fails_413():
    headers = get_authenticated_headers()
    # 16 MB dummy content
    oversized = b"x" * (16 * 1024 * 1024)
    files = {"file": ("large_resume.pdf", io.BytesIO(oversized), "application/pdf")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 413
    assert "large" in res.json()["detail"].lower()


def test_resume_upload_valid_docx_pipeline():
    headers = get_authenticated_headers()
    
    # Create sample docx in memory
    doc = docx.Document()
    doc.add_heading("Vinay Nalavade - Full Stack Engineer", level=1)
    doc.add_paragraph("Email: vinay@example.com | Phone: +91 9876543210")
    doc.add_paragraph("Summary: Experienced Full Stack Software Engineer with 4 years in Python, FastAPI, React.js, PostgreSQL, Docker, AWS, and CI/CD pipelines.")
    doc.add_paragraph("Experience: Developed scalable REST APIs improving response time by 40%. Engineered microservices serving 50k users.")
    doc.add_paragraph("Education: Bachelor of Technology (B.Tech) in Computer Science, 2021.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    files = {"file": ("Vinay_FullStack_Resume.docx", buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 200
    data = res.json()
    assert data["id"] > 0
    assert data["file_name"] == "Vinay_FullStack_Resume.docx"
    assert data["quality_score"] >= 40
    assert data["career_domain"] is not None
    assert len(data["strengths"]) > 0


def test_resume_upload_devops_cloud_engineer():
    headers = get_authenticated_headers()
    
    resume_text = """
    Arun Kumar - DevOps & Cloud Infrastructure Engineer
    Email: arun.devops@example.com | Phone: +91 9123456780
    Summary: DevOps specialist skilled in Kubernetes, Docker, Terraform, AWS, Jenkins, GitHub Actions, Linux, CI/CD pipelines, and Prometheus.
    Experience: Architected multi-region Kubernetes clusters on AWS with 99.99% uptime. Automated deployment pipelines reducing release time by 60%.
    Education: Bachelor of Engineering (B.E.) in Information Technology, 2020.
    """

    files = {"file": ("Arun_DevOps.txt", io.BytesIO(resume_text.encode("utf-8")), "text/plain")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 200
    data = res.json()
    assert "DevOps" in data["career_domain"] or "Cloud" in data["career_domain"] or "Infrastructure" in data["career_domain"] or data["quality_score"] > 50


def test_resume_upload_business_analyst():
    headers = get_authenticated_headers()
    
    resume_text = """
    Priya Sharma - Business Analyst & Product Specialist
    Email: priya.ba@example.com | Phone: +91 9988776655
    Summary: Business Analyst with expertise in Requirements Gathering, BRD/FRD creation, User Stories, Agile Scrum, Jira, Confluence, and Stakeholder Management.
    Experience: Spearheaded requirement elicitation workshops for fintech client. Delivered 12 sprint backlogs and increased delivery speed by 25%.
    Education: Master of Business Administration (MBA), 2022.
    """

    files = {"file": ("Priya_BA.txt", io.BytesIO(resume_text.encode("utf-8")), "text/plain")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 200
    data = res.json()
    assert "Business" in data["career_domain"] or "Product" in data["career_domain"] or data["quality_score"] > 50


def test_resume_upload_qa_automation_engineer():
    headers = get_authenticated_headers()
    
    resume_text = """
    Neha Patel - QA Automation Engineer
    Email: neha.qa@example.com | Phone: +91 9811223344
    Summary: Quality Assurance Engineer with 3+ years in Selenium WebDriver, Playwright, Test Automation, Postman, TestNG, Pytest, and Defect Management.
    Experience: Automated 450+ regression test cases reducing manual testing cycle by 70%. Executed API tests using Postman and Newman.
    Education: Bachelor of Computer Applications (BCA), 2021.
    """

    files = {"file": ("Neha_QA.txt", io.BytesIO(resume_text.encode("utf-8")), "text/plain")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 200
    data = res.json()
    assert "Quality" in data["career_domain"] or "QA" in data["career_domain"] or "Test" in data["career_domain"] or data["quality_score"] > 50


def test_resume_upload_general_technical_profile():
    headers = get_authenticated_headers()
    
    resume_text = """
    Technical Consultant Profile
    Email: consultant@example.com | Phone: +91 9777888999
    General technical overview, problem solving, analysis, reporting, and team collaboration.
    """

    files = {"file": ("General.txt", io.BytesIO(resume_text.encode("utf-8")), "text/plain")}
    res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert res.status_code == 200
    data = res.json()
    assert data["career_domain"] is not None


def test_view_uploaded_resume_file_authenticated():
    headers = get_authenticated_headers()
    
    # Upload first
    resume_text = "Sample resume content for download test"
    files = {"file": ("download_test.txt", io.BytesIO(resume_text.encode("utf-8")), "text/plain")}
    upload_res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert upload_res.status_code == 200
    resume_id = upload_res.json()["id"]

    # View file
    view_res = client.get(f"/api/v1/resumes/{resume_id}/file", headers=headers)
    assert view_res.status_code == 200
    assert b"Sample resume content" in view_res.content


def test_cors_preflight_and_post_resume_upload_github_pages():
    # OPTIONS
    options_headers = {
        "Origin": "https://vinaynalavade.github.io",
        "Access-Control-Request-Method": "POST",
        "Access-Control-Request-Headers": "authorization,content-type"
    }
    opt_res = client.options("/api/v1/resumes/upload", headers=options_headers)
    assert opt_res.status_code == 200
    assert opt_res.headers.get("access-control-allow-origin") == "https://vinaynalavade.github.io"
    assert opt_res.headers.get("access-control-allow-credentials") == "true"

    # POST
    headers = get_authenticated_headers()
    headers["Origin"] = "https://vinaynalavade.github.io"
    files = {"file": ("cors_resume.txt", io.BytesIO(b"Candidate CORS Test Resume Content"), "text/plain")}
    post_res = client.post("/api/v1/resumes/upload", files=files, headers=headers)
    assert post_res.status_code == 200
    assert post_res.headers.get("access-control-allow-origin") == "https://vinaynalavade.github.io"
    assert post_res.headers.get("access-control-allow-credentials") == "true"
